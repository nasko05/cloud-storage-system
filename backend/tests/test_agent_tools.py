"""Read-only tools, content extraction, and proposal building."""

import io

import pytest
from conftest import auth_header, register_and_login, upload_file
from docx import Document

from app.agent import tools
from app.agent.extract import MAX_TEXT_OUTPUT, extract_text
from app.database import SessionLocal
from app.errors import ApiError


def _folder(client, token, name, parent="root"):
    return client.post(
        f"/v2/folders/{parent}/folders", json={"name": name}, headers=auth_header(token)
    ).json()["folderId"]


def _make_pdf(text: str) -> bytes:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
    ]
    stream = b"BT /F1 24 Tf 72 720 Td (" + text.encode("latin-1") + b") Tj ET"
    objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n"
    xref_pos = len(pdf)
    pdf += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += (
        b"trailer\n<< /Size " + str(len(objects) + 1).encode() + b" /Root 1 0 R >>\n"
        b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return bytes(pdf)


def _make_docx(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- list_tree --------------------------------------------------------------

def test_list_tree_renders_ids_paths_types(client):
    token, user_id, _ = register_and_login(client)
    docs = _folder(client, token, "Documents")
    upload_file(client, token, "a.txt", parent=docs)
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "list_tree", {})
    assert "/Documents/  [folder id=" in out
    assert "/Documents/a.txt  [file id=" in out
    assert "type=text/plain" in out


def test_list_tree_empty_drive(client):
    _, user_id, _ = register_and_login(client)
    with SessionLocal() as db:
        assert tools.execute_readonly(db, user_id, "list_tree", {}) == "(the drive is empty)"


def test_list_tree_truncates(client, monkeypatch):
    token, user_id, _ = register_and_login(client)
    for i in range(5):
        _folder(client, token, f"Folder{i}")
    monkeypatch.setattr(tools, "MAX_TREE_NODES", 2)
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "list_tree", {})
    assert "tree truncated" in out


def test_unknown_readonly_tool_is_bad_request(client):
    _, user_id, _ = register_and_login(client)
    with SessionLocal() as db, pytest.raises(ApiError):
        tools.execute_readonly(db, user_id, "frobnicate", {})


# --- read_file --------------------------------------------------------------

def test_read_file_returns_text(client):
    token, user_id, _ = register_and_login(client)
    fid = upload_file(client, token, "note.txt", content=b"hello there body text")
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "read_file", {"id": fid})
    assert "hello there body text" in out


def test_read_file_pending_upload_has_no_content(client):
    token, user_id, _ = register_and_login(client)
    init = client.post(
        "/v2/files/uploads", json={"filename": "p.txt", "size": 1}, headers=auth_header(token)
    ).json()
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "read_file", {"id": init["fileId"]})
    assert "no stored content" in out


def test_read_file_requires_id(client):
    _, user_id, _ = register_and_login(client)
    with SessionLocal() as db, pytest.raises(ApiError):
        tools.execute_readonly(db, user_id, "read_file", {})


def test_read_file_foreign_owner_denied(client):
    token_a, _, _ = register_and_login(client)
    fid = upload_file(client, token_a, "secret.txt", content=b"private")
    _, user_b_id, _ = register_and_login(client)
    with SessionLocal() as db, pytest.raises(ApiError):
        tools.execute_readonly(db, user_b_id, "read_file", {"id": fid})


# --- extract_text (unit) ----------------------------------------------------

def test_extract_text_decodes_textual_by_extension():
    assert "x = 1" in extract_text("script.py", "application/octet-stream", b"x = 1\n")


def test_extract_text_binary_descriptor():
    out = extract_text("pic.png", "image/png", b"\x89PNG\r\n\x00\x01\x02")
    assert "binary file" in out and "image/png" in out


def test_extract_text_truncates_large_text():
    out = extract_text("big.txt", "text/plain", b"a" * (MAX_TEXT_OUTPUT + 500))
    assert "file truncated" in out
    assert len(out) < MAX_TEXT_OUTPUT + 500


def test_extract_text_pdf():
    out = extract_text("invoice.pdf", "application/pdf", _make_pdf("Hello Invoice 2026"))
    assert "Hello Invoice 2026" in out


def test_extract_text_pdf_corrupt():
    out = extract_text("broken.pdf", "application/pdf", b"%PDF-1.4 not really a pdf")
    assert "could not extract" in out.lower() or "no extractable" in out.lower()


def test_extract_text_docx():
    out = extract_text("report.docx", "application/octet-stream", _make_docx("Quarterly budget report"))
    assert "Quarterly budget report" in out


def test_extract_text_docx_corrupt():
    out = extract_text("broken.docx", "application/octet-stream", b"PK\x03\x04 not a real docx")
    assert "could not extract" in out.lower()


def test_extract_text_pdf_no_text():
    out = extract_text("scan.pdf", "application/pdf", _make_pdf(""))
    assert "no extractable text" in out.lower()


def test_extract_text_docx_no_text():
    buffer = io.BytesIO()
    Document().save(buffer)  # a document with no paragraphs
    out = extract_text("empty.docx", "application/octet-stream", buffer.getvalue())
    assert "no extractable text" in out.lower()


# --- search_files -----------------------------------------------------------

def test_search_files_finds_by_name(client):
    token, user_id, _ = register_and_login(client)
    docs = _folder(client, token, "Documents")
    upload_file(client, token, "tax-return-2026.pdf", parent=docs)
    upload_file(client, token, "holiday.jpg")
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "search_files", {"query": "tax"})
    assert "/Documents/tax-return-2026.pdf" in out
    assert "holiday.jpg" not in out


def test_search_files_no_match(client):
    token, user_id, _ = register_and_login(client)
    upload_file(client, token, "a.txt")
    with SessionLocal() as db:
        out = tools.execute_readonly(db, user_id, "search_files", {"query": "zzznope"})
    assert "No files match" in out


def test_search_files_requires_query(client):
    _, user_id, _ = register_and_login(client)
    with SessionLocal() as db, pytest.raises(ApiError):
        tools.execute_readonly(db, user_id, "search_files", {"query": "   "})


# --- proposals --------------------------------------------------------------

def _call(name: str, arguments: str):
    return {"id": "c1", "type": "function", "function": {"name": name, "arguments": arguments}}


def test_is_mutating_classification():
    for name in ("create_folder", "move_node", "rename_node"):
        assert tools.is_mutating(name)
    for name in ("list_tree", "read_file", "search_files"):
        assert not tools.is_mutating(name)


def test_build_pending_summaries():
    create = tools.build_pending(_call("create_folder", '{"parentPath":"/","name":"Photos"}'))
    assert create["tool"] == "create_folder"
    assert "Create folder" in create["summary"] and "Photos" in create["summary"]

    move = tools.build_pending(_call("move_node", '{"nodeType":"file","id":"x","path":"/a.txt","destinationPath":"/Photos"}'))
    assert "Move /a.txt → /Photos" in move["summary"]

    rename = tools.build_pending(_call("rename_node", '{"nodeType":"file","id":"x","path":"/a.txt","newName":"b.txt"}'))
    assert "Rename /a.txt → b.txt" in rename["summary"]


def test_build_pending_tolerates_empty_args_and_unknown_tool():
    empty = tools.build_pending(_call("create_folder", ""))
    assert empty["args"] == {}
    other = tools.build_pending(_call("frobnicate", "{}"))
    assert "frobnicate" in other["summary"]


def test_build_pending_rejects_invalid_json():
    with pytest.raises(ApiError):
        tools.build_pending(_call("create_folder", "{not json"))


def test_parse_args_handles_dict_and_non_object():
    assert tools.parse_args({"a": 1}) == {"a": 1}
    assert tools.parse_args("[1,2,3]") == {}
    assert tools.parse_args(None) == {}
